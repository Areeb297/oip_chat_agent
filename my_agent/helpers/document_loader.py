"""Document loaders for various file formats"""
from pathlib import Path
from typing import List, Optional
from ..models import Document, DocumentType


class DocumentLoader:
    """Load and extract text from various document formats"""

    @staticmethod
    def load_pdf(file_path: str) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF required: pip install PyMuPDF")

        doc = fitz.open(file_path)
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")

        doc.close()
        return "\n\n".join(text_parts)

    @staticmethod
    def load_docx(file_path: str) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        doc = DocxDocument(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    @staticmethod
    def load_txt(file_path: str) -> str:
        """Load text from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            File content
        """
        return Path(file_path).read_text(encoding="utf-8")

    @classmethod
    def load(cls, file_path: str) -> Document:
        """Load document based on file extension.

        Args:
            file_path: Path to document

        Returns:
            Document object with content and metadata
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        if ext == ".pdf":
            content = cls.load_pdf(str(path))
            doc_type = DocumentType.PDF
        elif ext in [".docx", ".doc"]:
            content = cls.load_docx(str(path))
            doc_type = DocumentType.DOCX
        elif ext == ".txt":
            content = cls.load_txt(str(path))
            doc_type = DocumentType.TXT
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        return Document(
            content=content,
            source=path.name,
            doc_type=doc_type,
        )

    @classmethod
    def load_directory(
        cls,
        directory: str,
        extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """Load all documents from a directory.

        Args:
            directory: Path to directory
            extensions: List of file extensions to include (default: pdf, docx, txt)

        Returns:
            List of Document objects
        """
        if extensions is None:
            extensions = [".pdf", ".docx", ".txt"]

        dir_path = Path(directory)
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        documents = []
        for ext in extensions:
            for file_path in dir_path.glob(f"*{ext}"):
                try:
                    doc = cls.load(str(file_path))
                    documents.append(doc)
                    print(f"Loaded: {file_path.name}")
                except Exception as e:
                    print(f"Error loading {file_path.name}: {e}")

        return documents
